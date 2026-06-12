from io import BytesIO

from docx import Document


def build_docx(
    study: dict[str, object],
    node: dict[str, object],
    rows: list[dict[str, object]],
) -> bytes:
    document = Document()
    document.add_heading(str(study["title"]), level=0)
    document.add_paragraph(f"Müşteri: {study['client']} · Tesis: {study['facility']}")
    document.add_heading(str(node["name"]), level=1)
    document.add_paragraph(f"Ekipman tipi: {node['equipment_type']}")
    document.add_paragraph(f"Tasarım niyeti: {node['design_intent']}")

    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = [
        "Kılavuz kelime",
        "Sapma",
        "Neden",
        "Sonuç",
        "Önlem",
        "Şiddet",
        "Olasılık",
        "Risk",
    ]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        values = [
            row["guideword"],
            row["deviation"],
            row["cause"],
            row["consequence"],
            row["safeguard"],
            row["severity"],
            row["likelihood"],
            row["risk"],
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)

    output = BytesIO()
    document.save(output)
    return output.getvalue()

